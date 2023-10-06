from docx import Document
import os
import pandas as pd
import subprocess

def get_git_root_path():
    try:
        root = subprocess.check_output(['git', 'rev-parse', '--show-toplevel']).strip().decode('utf-8')
    except subprocess.CalledProcessError:
        print("Error: Not a git repository")
        return None
    return root

def export_strings_to_word():
    git_root = get_git_root_path()
    if git_root:
        excel_dir = os.path.normpath(os.path.join(git_root, 'OriginTable'))
        word_dir = os.path.normpath(os.path.join(git_root, 'OriginTable', 'StringText'))

        if not os.path.exists(word_dir):
            os.makedirs(word_dir)

        doc = Document()
        doc.add_heading('游戏文本', 0)

        unique_chars = set()

        for filename in os.listdir(excel_dir):
            if filename.endswith('.xlsx') or filename.endswith('.xls'):
                print(f"Processing file: {filename}")

                excel_path = os.path.join(excel_dir, filename)
                xls = pd.ExcelFile(excel_path)

                for sheet_name in xls.sheet_names:
                    print(f"  Processing sheet: {sheet_name}")

                    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
                    print(f"    Second row: {df.iloc[1].to_dict()}")

                    for col_name in df.columns:
                        if "&" in str(df[col_name][1]):
                            print(f"    Found '&' in column: {col_name}")

                            for cell in df[col_name][3:]:
                                if pd.notna(cell):
                                    unique_chars.update(str(cell))

        # 将所有不重复的字符添加到Word文档中
        doc.add_paragraph(''.join(unique_chars))

        # 保存Word文档
        doc.save(os.path.join(word_dir, 'UniqueGameText.docx'))
        print("Word document has been saved.")

    else:
        print("Not a git repository.")  

if __name__ == "__main__":
    export_strings_to_word()